"""
Document loader for Layer 1 - handles PDF, EPUB, and text files.
"""

import logging
from pathlib import Path
from typing import Dict, Tuple, Optional
import re

from PyPDF2 import PdfReader
import ebooklib
from ebooklib import epub
from docx import Document as DocxDocument

from ...models.document import Document, DocumentMetadata

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Layer 1: Document loading with precise location mapping.
    
    Loads documents and creates detailed mappings for copyright compliance.
    Never sends original content externally - only provides data to Layer 2.
    """
    
    SUPPORTED_FORMATS = {".pdf", ".txt", ".epub", ".docx"}
    
    def _detect_actual_file_type(self, file_path: Path) -> str:
        """
        Detect the actual file type by examining file content, not just extension.
        
        Returns:
            Detected file type ("pdf", "txt", "docx", "epub", "unknown")
        """
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            # PDF files start with %PDF
            if header.startswith(b'%PDF'):
                return "pdf"
            
            # DOCX files are ZIP archives with specific structure
            if header.startswith(b'PK'):  # ZIP/DOCX signature
                try:
                    # Try to open as DOCX to confirm
                    DocxDocument(str(file_path))
                    return "docx"
                except:
                    pass
                
                # Could be EPUB (also ZIP-based)
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read(100)
                    if b'mimetype' in content and b'epub' in content:
                        return "epub"
                except:
                    pass
            
            # Try to decode as text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    f.read(100)  # Try to read some text
                return "txt"
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        f.read(100)
                    return "txt"  # Might be text with different encoding
                except:
                    pass
            
            return "unknown"
            
        except Exception as e:
            logger.warning(f"Could not detect file type for {file_path}: {e}")
            return "unknown"
    
    def _validate_file_type(self, file_path: Path) -> Tuple[str, bool]:
        """
        Validate that file extension matches actual content.
        
        Returns:
            (detected_type, matches_extension)
        """
        expected_type = file_path.suffix.lower().lstrip('.')
        detected_type = self._detect_actual_file_type(file_path)
        matches = (expected_type == detected_type)
        
        if not matches:
            logger.warning(
                f"File type mismatch: {file_path.name} "
                f"has extension '.{expected_type}' but appears to be '{detected_type}'"
            )
        
        return detected_type, matches
    
    def load(self, file_path: Path) -> Document:
        """
        Load a document and create location mappings.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Document with content and precise location mappings
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
            
        if file_path.suffix.lower() not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {file_path.suffix}")
        
        logger.info(f"Loading document: {file_path}")
        
        # Validate file type matches extension
        detected_type, matches_extension = self._validate_file_type(file_path)
        
        if not matches_extension:
            logger.warning(
                f"File extension mismatch detected! "
                f"Using detected type '{detected_type}' instead of extension '{file_path.suffix}'"
            )
        
        # Extract content based on DETECTED file type (not extension)
        if detected_type == "pdf":
            return self._load_pdf(file_path)
        elif detected_type == "txt":
            return self._load_text(file_path)
        elif detected_type == "epub":
            return self._load_epub(file_path)
        elif detected_type == "docx":
            return self._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported or undetected file type: {detected_type} for {file_path}")
    
    def _load_pdf(self, file_path: Path) -> Document:
        """Load PDF document with page-level mapping."""
        try:
            reader = PdfReader(str(file_path))
            
            text_content = ""
            page_mapping = {}
            paragraph_mapping = {}
            paragraph_counter = 0
            
            for page_num, page in enumerate(reader.pages, 1):
                page_start = len(text_content)
                page_text = page.extract_text()
                
                # Split into paragraphs (double newline or significant whitespace)
                paragraphs = re.split(r'\n\s*\n', page_text)
                
                for para_text in paragraphs:
                    if para_text.strip():  # Skip empty paragraphs
                        para_start = len(text_content)
                        text_content += para_text.strip() + "\n\n"
                        para_end = len(text_content)
                        
                        paragraph_mapping[paragraph_counter] = (para_start, para_end)
                        paragraph_counter += 1
                
                page_end = len(text_content)
                page_mapping[page_num] = (page_start, page_end)
            
            # Create metadata
            metadata = DocumentMetadata(
                file_path=file_path,
                file_size=file_path.stat().st_size,
                file_format="pdf",
                page_count=len(reader.pages),
                word_count=len(text_content.split())
            )
            
            # Try to extract title from PDF metadata
            if reader.metadata:
                metadata.title = reader.metadata.get('/Title')
                metadata.author = reader.metadata.get('/Author')
            
            return Document(
                metadata=metadata,
                text_content=text_content,
                page_mapping=page_mapping,
                paragraph_mapping=paragraph_mapping,
                section_mapping={}  # Could be enhanced with section detection
            )
            
        except Exception as e:
            logger.error(f"Failed to load PDF {file_path}: {e}")
            raise
    
    def _load_text(self, file_path: Path) -> Document:
        """Load plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Create paragraph mapping
            paragraphs = re.split(r'\n\s*\n', content)
            paragraph_mapping = {}
            current_pos = 0
            
            for i, para in enumerate(paragraphs):
                if para.strip():
                    start_pos = content.find(para, current_pos)
                    end_pos = start_pos + len(para)
                    paragraph_mapping[i] = (start_pos, end_pos)
                    current_pos = end_pos
            
            metadata = DocumentMetadata(
                file_path=file_path,
                file_size=file_path.stat().st_size,
                file_format="txt",
                word_count=len(content.split())
            )
            
            return Document(
                metadata=metadata,
                text_content=content,
                page_mapping={1: (0, len(content))},  # Single "page"
                paragraph_mapping=paragraph_mapping,
                section_mapping={}
            )
            
        except Exception as e:
            logger.error(f"Failed to load text file {file_path}: {e}")
            raise
    
    def _load_epub(self, file_path: Path) -> Document:
        """Load EPUB file with chapter mapping."""
        try:
            book = epub.read_epub(str(file_path))
            
            text_content = ""
            page_mapping = {}  # Will map to chapters
            paragraph_mapping = {}
            section_mapping = {}
            paragraph_counter = 0
            
            chapter_num = 1
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    chapter_start = len(text_content)
                    
                    # Extract text from HTML
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(item.content, 'html.parser')
                    chapter_text = soup.get_text()
                    
                    # Process paragraphs
                    paragraphs = re.split(r'\n\s*\n', chapter_text)
                    for para_text in paragraphs:
                        if para_text.strip():
                            para_start = len(text_content)
                            text_content += para_text.strip() + "\n\n"
                            para_end = len(text_content)
                            paragraph_mapping[paragraph_counter] = (para_start, para_end)
                            paragraph_counter += 1
                    
                    chapter_end = len(text_content)
                    page_mapping[chapter_num] = (chapter_start, chapter_end)
                    
                    # Try to extract chapter title
                    title_tag = soup.find(['h1', 'h2', 'h3'])
                    if title_tag:
                        section_mapping[title_tag.get_text().strip()] = (chapter_start, chapter_end)
                    
                    chapter_num += 1
            
            metadata = DocumentMetadata(
                file_path=file_path,
                file_size=file_path.stat().st_size,
                file_format="epub",
                page_count=chapter_num - 1,
                word_count=len(text_content.split()),
                title=book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else None,
                author=book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else None
            )
            
            return Document(
                metadata=metadata,
                text_content=text_content,
                page_mapping=page_mapping,
                paragraph_mapping=paragraph_mapping,
                section_mapping=section_mapping
            )
            
        except Exception as e:
            logger.error(f"Failed to load EPUB {file_path}: {e}")
            raise
    
    def _load_docx(self, file_path: Path) -> Document:
        """Load Word document."""
        try:
            doc = DocxDocument(str(file_path))
            
            text_content = ""
            paragraph_mapping = {}
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_start = len(text_content)
                    text_content += paragraph.text + "\n\n"
                    para_end = len(text_content)
                    paragraph_mapping[i] = (para_start, para_end)
            
            metadata = DocumentMetadata(
                file_path=file_path,
                file_size=file_path.stat().st_size,
                file_format="docx",
                word_count=len(text_content.split())
            )
            
            # Try to extract metadata from document properties
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata.title = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata.author = doc.core_properties.author
            
            return Document(
                metadata=metadata,
                text_content=text_content,
                page_mapping={1: (0, len(text_content))},  # Single "page"
                paragraph_mapping=paragraph_mapping,
                section_mapping={}
            )
            
        except Exception as e:
            logger.error(f"Failed to load DOCX {file_path}: {e}")
            raise
    
    # Public methods for testing and utilities
    def _is_supported_format(self, file_path: Path) -> bool:
        """Check if file extension is supported."""
        return file_path.suffix.lower() in self.SUPPORTED_FORMATS
    
    def _get_file_format(self, file_path: Path) -> str:
        """Get file format from extension."""
        return file_path.suffix.lower().lstrip('.')
    
    def get_file_type_info(self, file_path: Path) -> Dict[str, str]:
        """
        Get detailed file type information for debugging/testing.
        
        Returns:
            Dictionary with extension, detected_type, and matches info
        """
        expected_type = file_path.suffix.lower().lstrip('.')
        detected_type, matches = self._validate_file_type(file_path)
        
        return {
            'extension': expected_type,
            'detected_type': detected_type,
            'matches': matches,
            'supported': self._is_supported_format(file_path)
        }