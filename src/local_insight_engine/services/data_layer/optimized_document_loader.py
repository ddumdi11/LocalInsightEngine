"""
Optimized document loader for Layer 1 - Performance-enhanced PDF processing.
LocalInsightEngine v0.1.1 - Layer 1: High-Performance Data Layer

Key Optimizations:
- Streaming PDF processing to reduce memory footprint
- Efficient string building with list accumulation
- Memory-aware chunking for large documents
- Async support for concurrent processing
- Better memory management and cleanup
"""

import asyncio
import logging
import gc
import io
from pathlib import Path
from typing import Dict, Tuple, Optional, List, Iterator, Union
from dataclasses import dataclass
from contextlib import contextmanager
import re

try:
    import pymupdf as fitz  # PyMuPDF - much faster than PyPDF2
    PYMUPDF_AVAILABLE = True
except ImportError:
    from PyPDF2 import PdfReader
    PYMUPDF_AVAILABLE = False

import ebooklib
from ebooklib import epub
from docx import Document as DocxDocument

from ...models.document import Document, DocumentMetadata


logger = logging.getLogger(__name__)


@dataclass
class ProcessingStats:
    """Statistics for document processing performance."""
    pages_processed: int = 0
    paragraphs_processed: int = 0
    total_chars: int = 0
    processing_time_seconds: float = 0.0
    peak_memory_mb: float = 0.0
    chunks_created: int = 0


class StreamingDocumentLoader:
    """
    High-performance document loader with streaming capabilities.

    Key features:
    - Memory-efficient processing with configurable chunk sizes
    - Support for PyMuPDF (faster) or PyPDF2 (fallback)
    - Streaming text extraction for large documents
    - Async processing support
    - Optimized memory management
    """

    SUPPORTED_FORMATS = {".pdf", ".txt", ".epub", ".docx"}
    DEFAULT_CHUNK_SIZE = 1024 * 1024  # 1MB chunks for large documents
    MEMORY_THRESHOLD_MB = 100  # Switch to streaming mode above this size

    def __init__(self,
                 chunk_size: int = DEFAULT_CHUNK_SIZE,
                 memory_threshold_mb: int = MEMORY_THRESHOLD_MB,
                 enable_streaming: bool = True):
        """
        Initialize the optimized document loader.

        Args:
            chunk_size: Size of text chunks for streaming processing
            memory_threshold_mb: File size threshold to enable streaming mode
            enable_streaming: Whether to use streaming for large files
        """
        self.chunk_size = chunk_size
        self.memory_threshold_mb = memory_threshold_mb
        self.enable_streaming = enable_streaming
        self.stats = ProcessingStats()

        # Log which PDF backend is being used
        if PYMUPDF_AVAILABLE:
            logger.info("Using PyMuPDF (fitz) for high-performance PDF processing")
        else:
            logger.warning("PyMuPDF not available, falling back to PyPDF2 (slower)")

    def _detect_actual_file_type(self, file_path: Path) -> str:
        """
        Detect the actual file type by examining file content efficiently.

        Returns:
            Detected file type ("pdf", "txt", "docx", "epub", "unknown")
        """
        try:
            # Read only the first 64 bytes for faster detection
            with open(file_path, 'rb') as f:
                header = f.read(64)

            # PDF files start with %PDF
            if header.startswith(b'%PDF'):
                return "pdf"

            # DOCX/EPUB files are ZIP archives
            if header.startswith(b'PK'):
                # Quick check for EPUB mimetype
                if b'mimetype' in header[:64]:
                    return "epub"

                # Try DOCX validation (lighter check)
                try:
                    with open(file_path, 'rb') as f:
                        # Look for DOCX-specific content types
                        content = f.read(1024)  # Read more for DOCX detection
                    if b'word' in content.lower() or b'document.xml' in content:
                        return "docx"
                except:
                    pass

                # Could still be EPUB with different structure
                return "epub"  # Default for ZIP-based files

            # Try to decode as text (with size limit for large files)
            max_text_check = min(1024, file_path.stat().st_size)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    f.read(max_text_check)
                return "txt"
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        f.read(max_text_check)
                    return "txt"
                except:
                    pass

            return "unknown"

        except Exception as e:
            logger.warning(f"Could not detect file type for {file_path}: {e}")
            return "unknown"

    def _validate_file_type(self, file_path: Path) -> Tuple[str, bool]:
        """
        Validate that file extension matches actual content.

        Returns:
            (detected_type, matches_extension)
        """
        expected_type = file_path.suffix.lower().lstrip('.')
        detected_type = self._detect_actual_file_type(file_path)
        matches = (expected_type == detected_type)

        if not matches:
            logger.warning(
                f"File type mismatch: {file_path.name} "
                f"has extension '.{expected_type}' but appears to be '{detected_type}'"
            )

        return detected_type, matches

    @contextmanager
    def _memory_management(self):
        """Context manager for memory cleanup during processing."""
        try:
            gc.collect()  # Clean up before processing
            yield
        finally:
            gc.collect()  # Clean up after processing

    def _should_use_streaming(self, file_path: Path) -> bool:
        """Determine if streaming mode should be used for this file."""
        if not self.enable_streaming:
            return False

        file_size_mb = file_path.stat().st_size / 1024 / 1024
        return file_size_mb > self.memory_threshold_mb

    def load(self, file_path: Path) -> Document:
        """
        Load a document with optimized performance.

        Args:
            file_path: Path to document file

        Returns:
            Document with content and precise location mappings
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        if file_path.suffix.lower() not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {file_path.suffix}")

        logger.info(f"Loading document with optimized loader: {file_path}")

        # Reset stats
        self.stats = ProcessingStats()

        # Validate file type matches extension
        detected_type, matches_extension = self._validate_file_type(file_path)

        if not matches_extension:
            logger.warning(
                f"File extension mismatch! Using detected type '{detected_type}'"
            )

        # Use memory management context
        with self._memory_management():
            # Extract content based on detected file type
            if detected_type == "pdf":
                return self._load_pdf_optimized(file_path)
            elif detected_type == "txt":
                return self._load_text_optimized(file_path)
            elif detected_type == "epub":
                return self._load_epub_optimized(file_path)
            elif detected_type == "docx":
                return self._load_docx_optimized(file_path)
            else:
                raise ValueError(f"Unsupported file type: {detected_type} for {file_path}")

    def _load_pdf_optimized(self, file_path: Path) -> Document:
        """Load PDF document with performance optimizations."""
        use_streaming = self._should_use_streaming(file_path)

        if PYMUPDF_AVAILABLE:
            return self._load_pdf_pymupdf(file_path, use_streaming)
        else:
            return self._load_pdf_pypdf2_optimized(file_path, use_streaming)

    def _load_pdf_pymupdf(self, file_path: Path, use_streaming: bool) -> Document:
        """Load PDF using PyMuPDF (fastest option)."""
        try:
            doc = fitz.open(str(file_path))

            if use_streaming:
                return self._stream_pdf_pymupdf(file_path, doc)
            else:
                return self._load_pdf_pymupdf_standard(file_path, doc)

        except Exception as e:
            logger.error(f"Failed to load PDF with PyMuPDF {file_path}: {e}")
            raise

    def _load_pdf_pymupdf_standard(self, file_path: Path, doc) -> Document:
        """Standard PyMuPDF processing for smaller PDFs."""
        # Use list accumulation for efficient string building
        text_parts = []
        page_mapping = {}
        paragraph_mapping = {}
        paragraph_counter = 0

        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_start = sum(len(part) for part in text_parts)

            # Extract text with layout preservation
            page_text = page.get_text("text")

            # Process paragraphs efficiently
            paragraphs = re.split(r'\n\s*\n', page_text)

            for para_text in paragraphs:
                if para_text.strip():
                    para_start = sum(len(part) for part in text_parts)
                    cleaned_para = para_text.strip() + "\n\n"
                    text_parts.append(cleaned_para)

                    para_end = sum(len(part) for part in text_parts)
                    paragraph_mapping[paragraph_counter] = (para_start, para_end)
                    paragraph_counter += 1

            page_end = sum(len(part) for part in text_parts)
            page_mapping[page_num + 1] = (page_start, page_end)

        # Join all parts at once (more efficient than incremental concatenation)
        text_content = ''.join(text_parts)

        # Update stats
        self.stats.pages_processed = doc.page_count
        self.stats.paragraphs_processed = paragraph_counter
        self.stats.total_chars = len(text_content)

        # Create metadata
        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=file_path.stat().st_size,
            file_format="pdf",
            page_count=doc.page_count,
            word_count=len(text_content.split())
        )

        # Extract metadata from PDF
        pdf_metadata = doc.metadata
        if pdf_metadata:
            metadata.title = pdf_metadata.get('title')
            metadata.author = pdf_metadata.get('author')

        doc.close()  # Important: close document to free memory

        return Document(
            metadata=metadata,
            text_content=text_content,
            page_mapping=page_mapping,
            paragraph_mapping=paragraph_mapping,
            section_mapping={}
        )

    def _stream_pdf_pymupdf(self, file_path: Path, doc) -> Document:
        """Stream-process large PDFs to minimize memory usage."""
        logger.info(f"Using streaming mode for large PDF: {file_path.name}")

        # Process in chunks to manage memory
        page_mapping = {}
        paragraph_mapping = {}
        paragraph_counter = 0
        total_length = 0

        # First pass: collect metadata and page boundaries
        page_boundaries = []

        # Process pages in chunks
        chunk_pages = 10  # Process 10 pages at a time
        text_chunks = []

        for chunk_start in range(0, doc.page_count, chunk_pages):
            chunk_end = min(chunk_start + chunk_pages, doc.page_count)
            chunk_text_parts = []

            for page_num in range(chunk_start, chunk_end):
                page = doc[page_num]
                page_start = total_length

                page_text = page.get_text("text")
                paragraphs = re.split(r'\n\s*\n', page_text)

                page_text_parts = []
                for para_text in paragraphs:
                    if para_text.strip():
                        para_start = total_length + sum(len(p) for p in page_text_parts)
                        cleaned_para = para_text.strip() + "\n\n"
                        page_text_parts.append(cleaned_para)

                        para_end = para_start + len(cleaned_para)
                        paragraph_mapping[paragraph_counter] = (para_start, para_end)
                        paragraph_counter += 1

                page_text_content = ''.join(page_text_parts)
                chunk_text_parts.append(page_text_content)

                page_end = total_length + len(page_text_content)
                page_mapping[page_num + 1] = (page_start, page_end)
                total_length = page_end

            # Store chunk and clean up
            chunk_content = ''.join(chunk_text_parts)
            text_chunks.append(chunk_content)

            # Memory cleanup
            if chunk_start % 50 == 0:  # Every 50 pages
                gc.collect()

        # Final assembly
        text_content = ''.join(text_chunks)

        # Update stats
        self.stats.pages_processed = doc.page_count
        self.stats.paragraphs_processed = paragraph_counter
        self.stats.total_chars = len(text_content)
        self.stats.chunks_created = len(text_chunks)

        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=file_path.stat().st_size,
            file_format="pdf",
            page_count=doc.page_count,
            word_count=len(text_content.split())
        )

        # Extract metadata
        pdf_metadata = doc.metadata
        if pdf_metadata:
            metadata.title = pdf_metadata.get('title')
            metadata.author = pdf_metadata.get('author')

        doc.close()

        return Document(
            metadata=metadata,
            text_content=text_content,
            page_mapping=page_mapping,
            paragraph_mapping=paragraph_mapping,
            section_mapping={}
        )

    def _load_pdf_pypdf2_optimized(self, file_path: Path, use_streaming: bool) -> Document:
        """Optimized PyPDF2 processing (fallback when PyMuPDF unavailable)."""
        try:
            reader = PdfReader(str(file_path))

            if use_streaming:
                return self._stream_pdf_pypdf2(file_path, reader)
            else:
                return self._load_pdf_pypdf2_standard(file_path, reader)

        except Exception as e:
            logger.error(f"Failed to load PDF with PyPDF2 {file_path}: {e}")
            raise

    def _load_pdf_pypdf2_standard(self, file_path: Path, reader) -> Document:
        """Standard optimized PyPDF2 processing."""
        # Use list accumulation instead of string concatenation
        text_parts = []
        page_mapping = {}
        paragraph_mapping = {}
        paragraph_counter = 0

        for page_num, page in enumerate(reader.pages, 1):
            page_start = sum(len(part) for part in text_parts)
            page_text = page.extract_text()

            # Split into paragraphs
            paragraphs = re.split(r'\n\s*\n', page_text)

            for para_text in paragraphs:
                if para_text.strip():
                    para_start = sum(len(part) for part in text_parts)
                    cleaned_para = para_text.strip() + "\n\n"
                    text_parts.append(cleaned_para)

                    para_end = sum(len(part) for part in text_parts)
                    paragraph_mapping[paragraph_counter] = (para_start, para_end)
                    paragraph_counter += 1

            page_end = sum(len(part) for part in text_parts)
            page_mapping[page_num] = (page_start, page_end)

        # Single join operation
        text_content = ''.join(text_parts)

        # Update stats
        self.stats.pages_processed = len(reader.pages)
        self.stats.paragraphs_processed = paragraph_counter
        self.stats.total_chars = len(text_content)

        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=file_path.stat().st_size,
            file_format="pdf",
            page_count=len(reader.pages),
            word_count=len(text_content.split())
        )

        if reader.metadata:
            metadata.title = reader.metadata.get('/Title')
            metadata.author = reader.metadata.get('/Author')

        return Document(
            metadata=metadata,
            text_content=text_content,
            page_mapping=page_mapping,
            paragraph_mapping=paragraph_mapping,
            section_mapping={}
        )

    def _stream_pdf_pypdf2(self, file_path: Path, reader) -> Document:
        """Memory-efficient streaming for large PDFs with PyPDF2."""
        logger.info(f"Using streaming mode for large PDF (PyPDF2): {file_path.name}")

        page_mapping = {}
        paragraph_mapping = {}
        paragraph_counter = 0
        text_chunks = []
        total_length = 0

        # Process pages in smaller chunks
        chunk_size = 5  # Smaller chunks for PyPDF2 due to memory constraints

        for chunk_start in range(0, len(reader.pages), chunk_size):
            chunk_end = min(chunk_start + chunk_size, len(reader.pages))
            chunk_text_parts = []

            for page_idx in range(chunk_start, chunk_end):
                page = reader.pages[page_idx]
                page_num = page_idx + 1

                page_start = total_length
                page_text = page.extract_text()

                paragraphs = re.split(r'\n\s*\n', page_text)
                page_text_parts = []

                for para_text in paragraphs:
                    if para_text.strip():
                        para_start = total_length + sum(len(p) for p in page_text_parts)
                        cleaned_para = para_text.strip() + "\n\n"
                        page_text_parts.append(cleaned_para)

                        para_end = para_start + len(cleaned_para)
                        paragraph_mapping[paragraph_counter] = (para_start, para_end)
                        paragraph_counter += 1

                page_content = ''.join(page_text_parts)
                chunk_text_parts.append(page_content)

                page_end = total_length + len(page_content)
                page_mapping[page_num] = (page_start, page_end)
                total_length = page_end

            chunk_content = ''.join(chunk_text_parts)
            text_chunks.append(chunk_content)

            # Frequent memory cleanup for PyPDF2
            gc.collect()

        text_content = ''.join(text_chunks)

        self.stats.pages_processed = len(reader.pages)
        self.stats.paragraphs_processed = paragraph_counter
        self.stats.total_chars = len(text_content)
        self.stats.chunks_created = len(text_chunks)

        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=file_path.stat().st_size,
            file_format="pdf",
            page_count=len(reader.pages),
            word_count=len(text_content.split())
        )

        if reader.metadata:
            metadata.title = reader.metadata.get('/Title')
            metadata.author = reader.metadata.get('/Author')

        return Document(
            metadata=metadata,
            text_content=text_content,
            page_mapping=page_mapping,
            paragraph_mapping=paragraph_mapping,
            section_mapping={}
        )

    def _load_text_optimized(self, file_path: Path) -> Document:
        """Optimized text file loading with streaming support."""
        try:
            file_size = file_path.stat().st_size

            if self._should_use_streaming(file_path):
                return self._stream_text_file(file_path)
            else:
                return self._load_text_standard(file_path)

        except Exception as e:
            logger.error(f"Failed to load text file {file_path}: {e}")
            raise

    def _load_text_standard(self, file_path: Path) -> Document:
        """Standard text file loading for smaller files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Efficient paragraph mapping
        paragraphs = re.split(r'\n\s*\n', content)
        paragraph_mapping = {}
        current_pos = 0

        for i, para in enumerate(paragraphs):
            if para.strip():
                start_pos = content.find(para, current_pos)
                end_pos = start_pos + len(para)
                paragraph_mapping[i] = (start_pos, end_pos)
                current_pos = end_pos

        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=file_path.stat().st_size,
            file_format="txt",
            word_count=len(content.split())
        )

        return Document(
            metadata=metadata,
            text_content=content,
            page_mapping={1: (0, len(content))},
            paragraph_mapping=paragraph_mapping,
            section_mapping={}
        )

    def _stream_text_file(self, file_path: Path) -> Document:
        """Stream large text files in chunks."""
        logger.info(f"Streaming large text file: {file_path.name}")

        text_parts = []
        paragraph_mapping = {}
        paragraph_counter = 0
        total_length = 0

        with open(file_path, 'r', encoding='utf-8') as f:
            while True:
                chunk = f.read(self.chunk_size)
                if not chunk:
                    break

                # Process paragraphs in chunk
                paragraphs = re.split(r'\n\s*\n', chunk)

                for para in paragraphs:
                    if para.strip():
                        para_start = total_length
                        text_parts.append(para.strip() + "\n\n")
                        para_end = total_length + len(text_parts[-1])

                        paragraph_mapping[paragraph_counter] = (para_start, para_end)
                        paragraph_counter += 1
                        total_length = para_end

        content = ''.join(text_parts)

        metadata = DocumentMetadata(
            file_path=file_path,
            file_size=file_path.stat().st_size,
            file_format="txt",
            word_count=len(content.split())
        )

        return Document(
            metadata=metadata,
            text_content=content,
            page_mapping={1: (0, len(content))},
            paragraph_mapping=paragraph_mapping,
            section_mapping={}
        )

    def _load_epub_optimized(self, file_path: Path) -> Document:
        """Optimized EPUB loading (same as original, already efficient)."""
        # EPUB processing is already fairly optimized in the original
        # Could be enhanced further, but focusing on PDF first
        try:
            book = epub.read_epub(str(file_path))

            text_parts = []
            page_mapping = {}
            paragraph_mapping = {}
            section_mapping = {}
            paragraph_counter = 0

            chapter_num = 1
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    chapter_start = sum(len(part) for part in text_parts)

                    # Extract text from HTML
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(item.content, 'html.parser')
                    chapter_text = soup.get_text()

                    # Process paragraphs
                    paragraphs = re.split(r'\n\s*\n', chapter_text)
                    chapter_parts = []

                    for para_text in paragraphs:
                        if para_text.strip():
                            para_start = chapter_start + sum(len(p) for p in chapter_parts)
                            cleaned_para = para_text.strip() + "\n\n"
                            chapter_parts.append(cleaned_para)

                            para_end = para_start + len(cleaned_para)
                            paragraph_mapping[paragraph_counter] = (para_start, para_end)
                            paragraph_counter += 1

                    chapter_content = ''.join(chapter_parts)
                    text_parts.append(chapter_content)

                    chapter_end = chapter_start + len(chapter_content)
                    page_mapping[chapter_num] = (chapter_start, chapter_end)

                    # Extract chapter title
                    title_tag = soup.find(['h1', 'h2', 'h3'])
                    if title_tag:
                        section_mapping[title_tag.get_text().strip()] = (chapter_start, chapter_end)

                    chapter_num += 1

            text_content = ''.join(text_parts)

            metadata = DocumentMetadata(
                file_path=file_path,
                file_size=file_path.stat().st_size,
                file_format="epub",
                page_count=chapter_num - 1,
                word_count=len(text_content.split()),
                title=book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else None,
                author=book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else None
            )

            return Document(
                metadata=metadata,
                text_content=text_content,
                page_mapping=page_mapping,
                paragraph_mapping=paragraph_mapping,
                section_mapping=section_mapping
            )

        except Exception as e:
            logger.error(f"Failed to load EPUB {file_path}: {e}")
            raise

    def _load_docx_optimized(self, file_path: Path) -> Document:
        """Optimized DOCX loading with list accumulation."""
        try:
            doc = DocxDocument(str(file_path))

            text_parts = []
            paragraph_mapping = {}

            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_start = sum(len(part) for part in text_parts)
                    para_content = paragraph.text + "\n\n"
                    text_parts.append(para_content)

                    para_end = sum(len(part) for part in text_parts)
                    paragraph_mapping[i] = (para_start, para_end)

            text_content = ''.join(text_parts)

            metadata = DocumentMetadata(
                file_path=file_path,
                file_size=file_path.stat().st_size,
                file_format="docx",
                word_count=len(text_content.split())
            )

            # Extract metadata from document properties
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata.title = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata.author = doc.core_properties.author

            return Document(
                metadata=metadata,
                text_content=text_content,
                page_mapping={1: (0, len(text_content))},
                paragraph_mapping=paragraph_mapping,
                section_mapping={}
            )

        except Exception as e:
            logger.error(f"Failed to load DOCX {file_path}: {e}")
            raise

    # Async support methods
    async def load_async(self, file_path: Path) -> Document:
        """Async version of document loading."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.load, file_path)

    async def load_multiple_async(self, file_paths: List[Path]) -> List[Document]:
        """Load multiple documents concurrently."""
        tasks = [self.load_async(path) for path in file_paths]
        return await asyncio.gather(*tasks)

    # Utility methods
    def get_processing_stats(self) -> ProcessingStats:
        """Get processing statistics from last load operation."""
        return self.stats

    def get_file_type_info(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """
        Get detailed file type information for debugging/testing.

        Returns:
            Dictionary with extension, detected_type, and other metadata
        """
        expected_type = file_path.suffix.lower().lstrip('.')
        detected_type, matches = self._validate_file_type(file_path)

        file_size_mb = file_path.stat().st_size / 1024 / 1024
        will_stream = self._should_use_streaming(file_path)

        return {
            'extension': expected_type,
            'detected_type': detected_type,
            'matches': matches,
            'supported': file_path.suffix.lower() in self.SUPPORTED_FORMATS,
            'file_size_mb': file_size_mb,
            'will_use_streaming': will_stream,
            'pdf_backend': 'pymupdf' if PYMUPDF_AVAILABLE else 'pypdf2'
        }